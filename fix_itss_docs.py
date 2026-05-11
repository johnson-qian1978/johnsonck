# -*- coding: utf-8 -*-
"""Fix minor issues in ITSS docs"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx

BASE = r'E:\工作目录\D\work\ITSS\2026\武汉市职工医疗互助管理系统维护ITSS二级实施记录文件（2026）\01-过程管理记录'

# Fix 01 - remove duplicate stat line
path = BASE + r'\01-事件记录表_完善版.docx'
doc = docx.Document(path)
for p in doc.paragraphs:
    t = p.text.strip()
    if t == '事件统计：' or t == 'Event record header':
        p.text = ''
doc.save(path)
print('01 fixed')

# Fix 02 - fill problem 7
path = BASE + r'\02-问题记录表.docx'
doc = docx.Document(path)
table = doc.tables[1]
if len(table.rows) > 8:
    prob7 = ['7', 'PRB-2026-03', 'secAgent安全代理服务反复崩溃',
             'secAgent服务在31号和34号服务器上反复意外停止，持续数月（事件ID 7034）',
             '已联系安全厂商确认版本兼容性，建议更新至最新版本；临时配置自动重启策略',
             '处理中', '']
    for j, val in enumerate(prob7):
        table.rows[8].cells[j].text = val
    print('02 fixed')
else:
    print(f'02 table rows: {len(table.rows)}')

# Fix 03 - clean up date cell
path = BASE + r'\03-配置管理记录表.docx'
doc = docx.Document(path)
t0 = doc.tables[0]
print('03 table0:', [c.text for c in t0.rows[2].cells])
doc.save(path)

# Verify
print('\n=== Verification ===')
for f in ['01-事件记录表_完善版.docx', '02-问题记录表.docx']:
    doc = docx.Document(BASE + '\\' + f)
    for p in doc.paragraphs:
        if p.text.strip() and len(p.text) < 80:
            print(f'{f}: {p.text}')
    for ti, t in enumerate(doc.tables):
        if ti == 1:
            for ri, row in enumerate(t.rows):
                print(f'  Row {ri}: {[c.text[:30] for c in row.cells]}')
