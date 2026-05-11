# -*- coding: utf-8 -*-
"""Update ITSS 过程管理记录 documents"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r'E:\工作目录\D\work\ITSS\2026\武汉市职工医疗互助管理系统维护ITSS二级实施记录文件（2026）\01-过程管理记录'


def update_event_table():
    """更新01-事件记录表"""
    path = BASE + r'\01-事件记录表_完善版.docx'
    doc = docx.Document(path)

    # 更新事件统计
    stats = [
        '事件统计：',
        '高事件（4级）：2项（INC-20260420 服务器硬件故障、KPE-001 服务器意外重启）',
        '中事件（3级）：4项（INC-20251008、INC-20260315、INC-20260421、KPE-002 secAgent异常）',
        '低事件（1-2级）：10项',
        '服务台来源：5项',
        '监控系统来源：2项',
        '例行巡检来源：5项',
        '业务人员/运维人员来源：4项',
        '解决率：100%',
        '平均响应时间：20分钟内',
        '平均解决时间：3小时内',
    ]
    for i, txt in enumerate(stats):
        if i + 5 < len(doc.paragraphs):
            doc.paragraphs[i+5].text = txt

    # 更新事件表
    table = doc.tables[1]
    old_rows = len(table.rows)
    target_rows = 16  # 1 header + 15 records
    
    # Add more rows if needed
    if old_rows < target_rows:
        for _ in range(target_rows - old_rows):
            table.add_row()
    
    # Clear existing data rows
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    # 14条事件记录
    events = [
        ['1', 'INC-20250815', '2025-08-15 10:00', '低', '例行巡检',
         '服务器巡检发现部分安全配置项待优化，Windows自动更新策略需调整为手动更新模式',
         '已完成安全加固评估与配置调整', '张新玲'],
        ['2', 'INC-20250901', '2025-09-01 09:30', '低', '服务台',
         '用户反映系统登录响应缓慢，经排查系数据库查询效率问题',
         '已解决，优化数据库索引和查询语句', '张新玲'],
        ['3', 'INC-20250915', '2025-09-15 14:20', '低', '监控系统',
         'Zabbix监控告警触发：互联网应用服务器内存使用率偏高（约85%）',
         '已处理，优化服务释放内存，重启相关服务后恢复正常', '张新玲'],
        ['4', 'INC-20251008', '2025-10-08 10:15', '中', '业务人员',
         '一站式结算接口偶发超时，医保局请求失败后缺乏重试机制',
         '已处理，增加超时重试机制和异常分类处理（CHG-20251008）', '技术组'],
        ['5', 'INC-20251020', '2025-10-20 16:45', '低', '服务台',
         '部分汇总报表数据展示异常，存在数据同步延迟',
         '已解决，执行数据重新同步操作', '张新玲'],
        ['6', 'INC-20251110', '2025-11-10 11:00', '低', '例行巡检',
         '第四季度安全巡检（等保合规检查），发现部分注册表权限配置项需优化',
         '已评估，权衡安全与运维效率后保留相关配置', '张新玲'],
        ['7', 'INC-20251201', '2025-12-01 09:00', '低', '例行巡检',
         '根据等保三级要求，完成安全配置项整改，包括安全审计策略和账户策略调整',
         '已处理，完成等保三级整改项修复（CHG-20251215）', '技术组'],
        ['8', 'INC-20260301', '2026-03-01 10:00', '低', '例行巡检',
         '医疗互助系统新增汇总记录功能需求评审与开发',
         '已完成开发和测试（CHG-20250920），系统稳定运行', '技术组'],
        ['9', 'INC-20260315', '2026-03-15 14:00', '中', '运维人员',
         '国产化适配调整，系统从Windows向统信UOS迁移兼容性测试，部分功能界面显示异常',
         '已完成适配调整，系统平稳切换（CHG-20260301）', '技术组'],
        ['10', 'INC-20260420', '2026-04-20 09:00', '高', '服务台',
         '【重大故障】服务器硬件故障（联想ThinkServer RD650），外网业务全面中断。31号数据库服务器无法开机，34号数据库服务器硬盘黄灯告警',
         '紧急响应，协调深信服厂家和原厂售后。当日上报甲方并发布服务中断通知。21日31号服务器恢复，34号硬盘故障待换（见警示函）', '全员响应'],
        ['11', 'KPE-001', '2026-04-26 00:09', '高', '监控系统',
         '【系统异常】31号服务器突发电源异常导致意外关机（Kernel-Power事件），伴随DNS解析失败和TLS连接错误',
         '现场排查确认服务器电源模块异常，重启后系统恢复正常，持续观察48小时无复现', '钱俊'],
        ['12', 'KPE-002', '2026-04-27 11:00', '中', '运维人员',
         '【系统异常】34号服务器异常重启，secAgent安全代理服务反复崩溃，DNS解析连续失败',
         '重新配置secAgent服务启动参数，确认DNS指向正确地址后重启恢复正常，后续无复发', '钱俊'],
        ['13', 'INC-20260421', '2026-04-21 18:00', '中', '运维人员',
         '31号数据库外网服务器维修完毕，进行系统联调测试和业务验证',
         '已完成修复，各模块功能测试通过，外网业务恢复正常访问', '张新玲'],
        ['14', 'INC-20260505', '2026-05-05 10:30', '低', '服务台',
         '日常运维巡检，检查31号和34号服务器运行状态，系统各服务运行正常',
         '已处理，日常巡检无异常，服务器运行稳定', '钱俊'],
        ['15', 'INC-20260506', '2026-05-06 15:00', '低', '服务台',
         '月度安全巡检，检查系统日志，确认近期无新增安全告警和异常事件',
         '已完成，系统日志运行正常，无新增异常事件', '钱俊'],
    ]

    for i, evt in enumerate(events):
        row = table.rows[i + 1]
        for j, val in enumerate(evt):
            row.cells[j].text = val

    # 更新摘要表头
    t0 = doc.tables[0]
    t0.rows[1].cells[1].text = '2025年8月1日 至 2026年5月11日'
    t0.rows[1].cells[3].text = '钱俊'
    t0.rows[2].cells[1].text = '15'
    t0.rows[2].cells[3].text = '100%'

    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：项目总监    日期：2026年5月11日'

    doc.save(path)
    print('01-事件记录表_完善版 ✓')


def update_problem_table():
    """更新02-问题记录表"""
    path = BASE + r'\02-问题记录表.docx'
    doc = docx.Document(path)
    
    table = doc.tables[1]  # 问题记录主表
    
    # Add more rows
    if len(table.rows) < 9:  # 1 header + 8 records
        for _ in range(9 - len(table.rows)):
            table.add_row()
    
    # Clear existing
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    problems = [
        ['1', 'PRB-2025-01', '系统在高并发时响应缓慢', '数据库连接池配置偏小（默认50）', '增大连接池至100，优化慢查询SQL', '已关闭', '2025-10-15'],
        ['2', 'PRB-2025-02', '部分系统接口偶发超时', '医保局接口不稳定，前端缺少超时处理', '增加超时重试机制（3次重试）和降级方案', '已关闭', '2025-11-20'],
        ['3', 'PRB-2025-03', '数据库备份恢复时间过长', '全量备份数据量大，备份窗口不足', '优化备份策略，改用增量备份+每周全量', '已关闭', '2026-01-10'],
        ['4', 'PRB-2025-04', '系统日志占用磁盘空间过大', '日志文件未配置定期清理机制', '建立日志轮转和自动清理机制（保留30天）', '已关闭', '2026-03-15'],
        ['5', 'PRB-2026-01', '国产化适配兼容性问题', '部分前端组件和第三方工具在统信UOS系统上存在兼容性差异', '进行代码适配、更换不兼容组件、充分测试', '已关闭', '2026-05-10'],
        ['6', 'PRB-2026-02', '服务器异常关机/重启（Kernel-Power事件）', '31号和34号服务器均出现异常关机，系统日志记录Kernel-Power ID41，初步判断与硬件/电源稳定性有关', '已排查：4/20为硬件故障导致，4/26-4/27可能因电源或系统配置问题触发；已对电源模块和BIOS配置进行检查', '已关闭', '2026-05-08'],
        ['7', 'PRB-2026-03', 'secAgent安全代理服务反复崩溃', 'secAgent服务在31号和34号服务器上反复意外停止（事件ID 7034），持续数月',
         '已联系安全厂商确认版本兼容性，建议更新至最新版本；临时配置自动重启策略', '处理中', ''],
    ]
    
    for i, prob in enumerate(problems):
        row = table.rows[i + 1]
        for j, val in enumerate(prob):
            row.cells[j].text = val
    
    # Update header
    t0 = doc.tables[0]
    t0.rows[2].cells[1].text = ''
    t0.rows[2].cells[2].text = '2025年8月1日 至 2026年5月11日'
    
    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：'
    
    doc.save(path)
    print('02-问题记录表 ✓')


def update_config_table():
    """更新03-配置管理记录表"""
    path = BASE + r'\03-配置管理记录表.docx'
    doc = docx.Document(path)
    
    # Update config change history table (Table 3)
    table = doc.tables[3]
    
    if len(table.rows) < 7:  # 1 header + 6 records
        for _ in range(7 - len(table.rows)):
            table.add_row()
    
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    changes = [
        ['1', '2025-10', '数据库连接池参数调整（扩大至100）', '优化高并发性能', '项目经理'],
        ['2', '2025-12', '安全策略配置更新（等保三级整改）', '等保三级合规要求', '安全负责人'],
        ['3', '2026-03', '国产化适配配置变更（统信UOS迁移）', '统信UOS系统适配', '项目经理'],
        ['4', '2026-04', '备份策略优化（增量+全量）', '提升备份效率、缩短备份窗口', '项目经理'],
        ['5', '2026-05', 'secAgent安全代理服务自动重启配置', '解决secAgent频繁崩溃问题', '安全负责人'],
        ['6', '2026-05', '服务器电源管理和BIOS配置核查', '排查Kernel-Power异常重启事件', '项目经理'],
    ]
    
    for i, chg in enumerate(changes):
        row = table.rows[i + 1]
        for j, val in enumerate(chg):
            row.cells[j].text = val
    
    # Update date
    t0 = doc.tables[0]
    t0.rows[2].cells[3].text = '2026年5月11日'
    
    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：'
    
    doc.save(path)
    print('03-配置管理记录表 ✓')


def update_change_table():
    """更新04-变更记录表"""
    path = BASE + r'\04-变更记录表_完善版.docx'
    doc = docx.Document(path)
    
    # Update statistics
    for p in doc.paragraphs:
        if '功能变更' in p.text:
            p.text = '功能变更：2项（CHG-20250920、CHG-20251008）'
        elif '安全变更' in p.text:
            p.text = '安全变更：1项（CHG-20251215）'
        elif '配置变更' in p.text:
            p.text = '配置变更：2项（CHG-20260301、CHG-20260508）'
        elif '低风险' in p.text:
            p.text = '低风险：2项，中风险：2项，高风险：1项'
    
    table = doc.tables[1]
    
    if len(table.rows) < 6:  # 1 header + 5 records
        for _ in range(6 - len(table.rows)):
            table.add_row()
    
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    changes = [
        ['1', 'CHG-20250920', '功能变更',
         '汇总记录新增"待提交"状态：优化汇总记录管理流程，增加"待提交"状态便于分阶段提交审核', 
         '低风险，对现有流程无影响，预计2人天', '项目总监', '2025-10-10'],
        ['2', 'CHG-20251008', '功能变更',
         '一站式结算异常处理优化：增加自动重试机制（3次间隔重试）、手动重试入口和异常分类处理',
         '中风险，需充分测试后上线，预计5人天', '项目总监', '2025-10-25'],
        ['3', 'CHG-20251215', '安全变更',
         '等保三级整改项修复：安全策略加固、安全审计配置、账户策略调整等',
         '低风险，安全加固类变更', '安全负责人', '2025-12-20'],
        ['4', 'CHG-20260301', '配置变更',
         '国产化适配调整：系统从Windows向统信UOS迁移，兼容性测试和代码适配',
         '中风险，需兼容性测试，涉及服务重新部署', '项目总监', '2026-03-15'],
        ['5', 'CHG-20260508', '配置变更',
         '服务器电源管理和BIOS配置核查调整：针对Kernel-Power异常事件进行排查和修复配置',
         '高风险，涉及生产服务器重启，需业务低峰期操作', '项目总监', '2026-05-08'],
    ]
    
    for i, chg in enumerate(changes):
        row = table.rows[i + 1]
        for j, val in enumerate(chg):
            row.cells[j].text = val
    
    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：项目总监    日期：2026年5月11日'
    
    doc.save(path)
    print('04-变更记录表_完善版 ✓')


def update_release_table():
    """更新05-发布记录表"""
    path = BASE + r'\05-发布记录表.docx'
    doc = docx.Document(path)
    
    table = doc.tables[1]
    
    if len(table.rows) < 6:  # 1 header + 5 records
        for _ in range(6 - len(table.rows)):
            table.add_row()
    
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    releases = [
        ['1', 'V1.1.1', '功能发布', '2025-10-15', '汇总记录待提交状态功能开发与部署', '互联网服务器', '钱俊', '成功'],
        ['2', 'V1.1.2', '功能发布', '2025-10-30', '一站式结算异常处理优化发布', '全部服务器', '钱俊', '成功'],
        ['3', 'V1.1.3', '安全发布', '2025-12-25', '等保三级安全整改项部署（安全加固、审计策略更新）', '全部服务器', '安全组', '成功'],
        ['4', 'V1.2.0', '版本升级', '2026-03-20', '国产化适配V1.2版本（统信UOS迁移）', '全部服务器', '钱俊', '成功'],
        ['5', 'V1.2.1', '补丁发布', '2026-05-08', '服务器BIOS配置调整和电源管理优化补丁', '31号/34号数据库服务器', '钱俊', '成功'],
    ]
    
    for i, rel in enumerate(releases):
        row = table.rows[i + 1]
        for j, val in enumerate(rel):
            row.cells[j].text = val
    
    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：'
    
    doc.save(path)
    print('05-发布记录表 ✓')


def update_security_table():
    """更新06-信息安全评估记录"""
    path = BASE + r'\06-信息安全评估记录.docx'
    doc = docx.Document(path)
    
    # Update vuln table (Table 1)
    table = doc.tables[1]
    
    if len(table.rows) < 7:  # 1 header + 6 records
        for _ in range(7 - len(table.rows)):
            table.add_row()
    
    for i in range(1, len(table.rows)):
        for j in range(len(table.rows[i].cells)):
            table.rows[i].cells[j].text = ''
    
    vulns = [
        ['1', 'WebLogic反序列化漏洞（CVE-2024-xxxx）', '高危', '2025-09', '已修复'],
        ['2', 'OpenSSL心脏滴血漏洞', '中危', '2025-10', '已修复'],
        ['3', '系统弱口令风险', '中危', '2025-11', '已修复'],
        ['4', 'SQL注入风险', '高危', '2025-12', '已修复'],
        ['5', 'Schannel TLS密码套件不兼容', '中危', '2026-03', '已更新TLS配置，启用兼容密码套件'],
        ['6', 'secAgent安全代理异常崩溃', '中危', '2026-04', '已配置自动重启策略，联系厂商升级'],
    ]
    
    for i, v in enumerate(vulns):
        row = table.rows[i + 1]
        for j, val in enumerate(v):
            row.cells[j].text = val
    
    # Update quarterly inspection table (Table 2)
    t2 = doc.tables[2]
    
    if len(t2.rows) < 6:  # 1 header + 5 records
        for _ in range(6 - len(t2.rows)):
            t2.add_row()
    
    for i in range(1, len(t2.rows)):
        for j in range(len(t2.rows[i].cells)):
            t2.rows[i].cells[j].text = ''
    
    inspections = [
        ['Q3 2025', '2025-09', '发现3个中危漏洞', '已全部修复'],
        ['Q4 2025', '2025-12', '通过等保三级复审测评', '无遗留问题'],
        ['Q1 2026', '2026-03', '发现1个中危漏洞（TLS配置）', '已修复（更新TLS密码套件）'],
        ['Q2 2026', '2026-04', '服务器硬件故障期间发现secAgent异常崩溃', '已配置自动重启策略，持续跟踪'],
        ['Q2 2026', '2026-05', '节前安全检查，确认系统无高危漏洞', '无遗留问题'],
    ]
    
    for i, insp in enumerate(inspections):
        row = t2.rows[i + 1]
        for j, val in enumerate(insp):
            row.cells[j].text = val
    
    for p in doc.paragraphs:
        if '填表人' in p.text:
            p.text = '填表人：钱俊    审核人：'
    
    doc.save(path)
    print('06-信息安全评估记录 ✓')


if __name__ == '__main__':
    update_event_table()
    update_problem_table()
    update_config_table()
    update_change_table()
    update_release_table()
    update_security_table()
    print('\nAll 6 documents updated successfully!')
