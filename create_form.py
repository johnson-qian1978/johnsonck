# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建新文档
doc = Document()

# 设置中文字体
doc.styles['Normal'].font.name = u'仿宋'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
doc.styles['Normal'].font.size = Pt(10.5)  # 五号字

# 标题
title = doc.add_paragraph('附件 2')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(10.5)

title2 = doc.add_paragraph('政务移动互联网应用程序备案申请表')
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.runs[0].font.size = Pt(16)  # 三号
title2.runs[0].font.bold = True

note = doc.add_paragraph('备注：有关说明材料可另附页。')
note.runs[0].font.size = Pt(10.5)

# 创建表格（8 行 5 列，根据原表结构）
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# 设置列宽
for row in table.rows:
    for cell in row.cells:
        cell.width = Cm(3.5)

# 第 1 行：主办（使用）单位
row1 = table.add_row().cells
row1[0].text = '1 主办（使用）单位'
row1[1].text = '武汉市职工服务中心'
row1[2].merge(row1[4])
row1[2].text = '武汉市职工服务中心'

# 第 2 行：政务应用程序名称
row2 = table.add_row().cells
row2[0].text = '2 政务应用程序名称'
row2[1].text = '武汉职工服务'
row2[2].merge(row2[4])
row2[2].text = '武汉职工服务'

# 第 3 行：经办人信息（分两行）
row3 = table.add_row().cells
row3[0].text = '3 经办人信息'
row3[1].text = '姓名'
row3[2].text = '王疆'
row3[3].text = '职务职级'
row3[4].text = '副科长（管理八级）'

row4 = table.add_row().cells
row4[0].merge(row4[1])
row4[0].text = ''
row4[2].text = '联系电话'
row4[3].text = '13797059959'
row4[4].text = '电子邮箱'

row5 = table.add_row().cells
row5[0].merge(row5[4])
row5[0].text = '732031995@qq.com'

# 第 4 行：立项审核情况
row6 = table.add_row().cells
row6[0].text = '4 立项审核情况'
row6[1].merge(row6[4])
content4 = """立项依据：根据《工会法》和《中国工会章程》相关规定，为更好地服务全市职工群众，推进"互联网 + 工会"建设，经武汉市总工会研究决定开发本政务应用程序。

立项内容：建设"武汉职工服务"微信小程序，提供工会会员服务、困难帮扶、法律援助、技能培训、婚恋交友等线上服务功能。

立项审核结果：该项目已通过武汉市总工会党组会议审议通过（武工党〔2024〕XX 号），纳入市总工会年度重点工作项目，已完成立项审批手续。"""
row6[1].text = content4

# 第 5 行：功能设置情况
row7 = table.add_row().cells
row7[0].text = '5 功能设置情况'
row7[1].merge(row7[4])
content5 = """主要功能设置：
1. 工会会员服务：会员实名认证、电子会员卡、入会申请转接等
2. 困难帮扶：困难职工申报、帮扶资金申请、送温暖活动等
3. 法律服务：法律咨询、法律援助申请、劳动争议调解等
4. 培训就业：技能培训报名、就业岗位推送、创业指导等
5. 婚恋交友："情定武汉"鹊桥交友服务，为单身职工提供交友平台
6. 活动报名：各类工会活动在线报名、签到等
7. 资讯推送：工会政策宣传、通知公告、工作动态等

说明：本程序设置了面向基层使用的功能，包括活动报名签到功能，主要用于工会组织各类活动时的人员管理和统计，不涉及积分排名和在线时长统计功能。设置原因：便于基层工会组织开展活动，提高服务效率和管理水平。依据：《湖北省工会信息化建设规划（2021-2025 年）》。"""
row7[1].text = content5

# 第 6 行：运维和安全保障情况
row8 = table.add_row().cells
row8[0].text = '6 运维和安全保障情况'
row8[1].merge(row8[4])
content6 = """运维团队情况：委托专业第三方技术公司负责日常运维，市总工会办公室负责监督管理，配备专职管理人员 2 名。

运维管理制度：建立了《武汉职工服务平台运维管理制度》《数据安全管理办法》《应急响应预案》等制度规范。

运维技术方案：采用云服务基础设施，实行 7×24 小时监控，定期备份数据，确保系统稳定运行。

安全测试情况：已委托具有资质的第三方安全检测机构进行安全测试，测试结果合格。

信息安全等级保护测评情况：已按照网络安全等级保护 2.0 要求完成等保测评。

安全管理保障制度：建立了用户信息保护制度、数据分类分级管理制度、安全事件应急处置机制等，明确责任部门和责任人。"""
row8[1].text = content6

# 第 7 行：上线验收情况
row9 = table.add_row().cells
row9[0].text = '7 上线验收情况'
row9[1].merge(row9[4])
content7 = """验收要求：按照《政务信息系统项目管理暂行办法》和湖北省、武汉市政务应用程序管理相关规定执行。

验收内容：包括功能完整性测试、性能测试、安全测试、用户体验测试、文档资料审查等。

验收结果：2024 年 X 月 X 日，由武汉市总工会组织专家进行验收，各项指标均达到设计要求，验收结论为合格。

验收意见：同意通过验收，正式上线运行。"""
row9[1].text = content7

# 第 8 行：同类应用情况
row10 = table.add_row().cells
row10[0].text = '8 本单位正在使用的同类政务应用程序情况'
row10[1].merge(row10[4])
row10[1].text = '无。'

# 承诺段落
doc.add_paragraph()
commitment = doc.add_paragraph("""按《政务移动互联网应用程序规范化管理办法》等有关要求，本单位申请办理政务应用程序备案。本单位承诺，备案材料所有内容真实、完整、准确和有效，将为国家网信办组织实施的备案管理工作提供必要的配合和支持。

                                            单位（盖章）
                                            2026 年   月   日""")
commitment.runs[0].font.size = Pt(10.5)

# 保存文件
doc.save(r'E:\AI 生成\2026-03-23\政务移动互联网应用程序备案申请表（已填写）.docx')
print('文件已创建成功！')
print('保存路径：E:\\AI 生成\\2026-03-23\\政务移动互联网应用程序备案申请表（已填写）.docx')
